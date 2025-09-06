# ========= Copyright 2023-2024 @ CAMEL-AI.org. All Rights Reserved. =========
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
# ========= Copyright 2023-2024 @ CAMEL-AI.org. All Rights Reserved. =========
import os
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple, Union

from camel.logger import get_logger
from camel.toolkits.base import BaseToolkit
from camel.toolkits.function_tool import FunctionTool
from camel.utils import MCPServer, dependencies_required

logger = get_logger(__name__)


@MCPServer()
class FileWriteToolkit(BaseToolkit):
    r"""A toolkit for creating, writing, and modifying text in files.

    This class provides cross-platform (macOS, Linux, Windows) support for
    writing to various file formats (Markdown, DOCX, PDF, and plaintext),
    replacing text in existing files, automatic filename uniquification to
    prevent overwrites, custom encoding and enhanced formatting options for
    specialized formats.
    """

    def __init__(
        self,
        working_directory: Optional[str] = None,
        timeout: Optional[float] = None,
        default_encoding: str = "utf-8",
    ) -> None:
        r"""Initialize the FileWriteToolkit.

        Args:
            working_directory (str, optional): The default directory for
                output files. If not provided, it will be determined by the
                `CAMEL_WORKDIR` environment variable (if set). If the
                environment variable is not set, it defaults to
                `camel_working_dir`.
            timeout (Optional[float]): The timeout for the toolkit.
                (default: :obj:`None`)
            default_encoding (str): Default character encoding for text
                operations. (default: :obj:`utf-8`)

        """
        super().__init__(timeout=timeout)
        if working_directory:
            self.working_directory = Path(working_directory).resolve()
        else:
            camel_workdir = os.environ.get("CAMEL_WORKDIR")
            if camel_workdir:
                self.working_directory = Path(camel_workdir).resolve()
            else:
                self.working_directory = Path("./camel_working_dir").resolve()
        self.working_directory.mkdir(parents=True, exist_ok=True)
        self.default_encoding = default_encoding

        logger.info(
            f"FileWriteToolkit initialized with output directory"
            f": {self.working_directory}, encoding: {default_encoding}"
        )

    def _resolve_filepath(self, file_path: str) -> Path:
        r"""Convert the given string path to a Path object.

        If the provided path is not absolute, it is made relative to the
        default output directory. The filename part is sanitized to replace
        spaces and special characters with underscores, ensuring safe usage
        in downstream processing.

        Args:
            file_path (str): The file path to resolve.

        Returns:
            Path: A fully resolved (absolute) and sanitized Path object.
        """
        path_obj = Path(file_path)
        if not path_obj.is_absolute():
            path_obj = self.working_directory / path_obj

        sanitized_filename = self._sanitize_filename(path_obj.name)
        path_obj = path_obj.parent / sanitized_filename
        return path_obj.resolve()

    def _sanitize_filename(self, filename: str) -> str:
        r"""Sanitize a filename by replacing any character that is not
        alphanumeric, a dot (.), hyphen (-), or underscore (_) with an
        underscore (_).

        Args:
            filename (str): The original filename which may contain spaces or
                special characters.

        Returns:
            str: The sanitized filename with disallowed characters replaced by
                underscores.
        """
        safe = re.sub(r'[^\w\-.]', '_', filename)
        return safe

    def _write_text_file(
        self, file_path: Path, content: str, encoding: str = "utf-8"
    ) -> None:
        r"""Write text content to a plaintext file.

        Args:
            file_path (Path): The target file path.
            content (str): The text content to write.
            encoding (str): Character encoding to use. (default: :obj:`utf-8`)
        """
        with file_path.open("w", encoding=encoding) as f:
            f.write(content)

    def _generate_unique_filename(self, file_path: Path) -> Path:
        r"""Generate a unique filename if the target file already exists.

        Args:
            file_path (Path): The original file path.

        Returns:
            Path: A unique file path that doesn't exist yet.
        """
        if not file_path.exists():
            return file_path

        # Generate unique filename with timestamp and counter
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = file_path.stem
        suffix = file_path.suffix
        parent = file_path.parent

        # First try with timestamp
        new_path = parent / f"{stem}_{timestamp}{suffix}"
        if not new_path.exists():
            return new_path

        # If timestamp version exists, add counter
        counter = 1
        while True:
            new_path = parent / f"{stem}_{timestamp}_{counter}{suffix}"
            if not new_path.exists():
                return new_path
            counter += 1

    def _write_docx_file(self, file_path: Path, content: str) -> None:
        r"""Write text content to a DOCX file with default formatting.

        Args:
            file_path (Path): The target file path.
            content (str): The text content to write.
        """
        import docx

        # Use default formatting values
        font_name = 'Calibri'
        font_size = 11
        line_spacing = 1.0

        document = docx.Document()
        style = document.styles['Normal']
        style.font.name = font_name
        style.font.size = docx.shared.Pt(font_size)
        style.paragraph_format.line_spacing = line_spacing

        # Split content into paragraphs and add them
        for para_text in content.split('\n'):
            para = document.add_paragraph(para_text)
            para.style = style

        document.save(str(file_path))

    @dependencies_required('reportlab')
    def _write_pdf_file(
        self,
        file_path: Path,
        title: str,
        content: Union[str, List[List[str]]],
        use_latex: bool = False,
    ) -> None:
        r"""Write text content to a PDF file with LaTeX and table support.

        Args:
            file_path (Path): The target file path.
            title (str): The document title.
            content (Union[str, List[List[str]]]): The content to write. Can
                be:
                - String: Supports Markdown-style tables and LaTeX math
                expressions
                - List[List[str]]: Table data as list of rows for direct table
                rendering
            use_latex (bool): Whether to use LaTeX for math rendering.
                (default: :obj:`False`)
        """
        if use_latex:
            from pylatex import (
                Command,
                Document,
                Math,
                Section,
            )
            from pylatex.utils import (
                NoEscape,
            )

            doc = Document(documentclass="article")
            doc.packages.append(Command('usepackage', 'amsmath'))
            with doc.create(Section('Generated Content')):
                # Handle different content types
                if isinstance(content, str):
                    content_lines = content.split('\n')
                else:
                    # Convert table data to LaTeX table format
                    content_lines = []
                    if content:
                        # Add table header
                        table_header = (
                            r'\begin{tabular}{' + 'l' * len(content[0]) + '}'
                        )
                        content_lines.append(table_header)
                        content_lines.append(r'\hline')
                        for row in content:
                            row_content = (
                                ' & '.join(str(cell) for cell in row) + r' \\'
                            )
                            content_lines.append(row_content)
                            content_lines.append(r'\hline')
                        content_lines.append(r'\end{tabular}')

                for line in content_lines:
                    stripped_line = line.strip()

                    # Skip empty lines
                    if not stripped_line:
                        continue

                    # Convert Markdown-like headers
                    if stripped_line.startswith('## '):
                        header = stripped_line[3:]
                        doc.append(NoEscape(r'\subsection*{%s}' % header))
                        continue
                    elif stripped_line.startswith('# '):
                        header = stripped_line[2:]
                        doc.append(NoEscape(r'\section*{%s}' % header))
                        continue
                    elif stripped_line.strip() == '---':
                        doc.append(NoEscape(r'\hrule'))
                        continue

                    # Detect standalone math expressions like $...$
                    if (
                        stripped_line.startswith('$')
                        and stripped_line.endswith('$')
                        and len(stripped_line) > 1
                    ):
                        math_data = stripped_line[1:-1]
                        doc.append(Math(data=math_data))
                    else:
                        doc.append(NoEscape(stripped_line))
                    doc.append(NoEscape(r'\par'))

                doc.generate_pdf(str(file_path), clean_tex=True)

            logger.info(f"Wrote PDF (with LaTeX) to {file_path}")

        else:
            try:
                from reportlab.lib import colors
                from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
                from reportlab.lib.pagesizes import A4
                from reportlab.lib.styles import (
                    ParagraphStyle,
                    getSampleStyleSheet,
                )
                from reportlab.platypus import (
                    Paragraph,
                    SimpleDocTemplate,
                    Spacer,
                )

                # Register Chinese fonts
                chinese_font = self._register_chinese_font()

                # Create PDF document
                doc = SimpleDocTemplate(
                    str(file_path),
                    pagesize=A4,
                    rightMargin=72,
                    leftMargin=72,
                    topMargin=72,
                    bottomMargin=18,
                )

                # Get styles with Chinese font support
                styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'CustomTitle',
                    parent=styles['Heading1'],
                    fontSize=18,
                    spaceAfter=30,
                    alignment=TA_CENTER,
                    textColor=colors.black,
                    fontName=chinese_font,
                )

                heading_style = ParagraphStyle(
                    'CustomHeading',
                    parent=styles['Heading2'],
                    fontSize=14,
                    spaceAfter=12,
                    spaceBefore=20,
                    textColor=colors.black,
                    fontName=chinese_font,
                )

                body_style = ParagraphStyle(
                    'CustomBody',
                    parent=styles['Normal'],
                    fontSize=11,
                    spaceAfter=12,
                    alignment=TA_JUSTIFY,
                    textColor=colors.black,
                    fontName=chinese_font,
                )

                # Build story (content elements)
                story = []

                # Add title
                if title:
                    story.append(Paragraph(title, title_style))
                    story.append(Spacer(1, 12))

                # Handle different content types
                if isinstance(content, list) and all(
                    isinstance(row, list) for row in content
                ):
                    # Content is a table (List[List[str]])
                    if content:
                        table = self._create_pdf_table(content)
                        story.append(table)
                else:
                    # Content is a string, process normally
                    content_str = str(content)
                    self._process_text_content(
                        story, content_str, heading_style, body_style
                    )

                # Build PDF
                doc.build(story)
            except Exception as e:
                logger.error(f"Error creating PDF: {e}")

    def _process_text_content(
        self, story, content: str, heading_style, body_style
    ):
        r"""Process text content and add to story.

        Args:
            story: The reportlab story list to append to
            content (str): The text content to process
            heading_style: Style for headings
            body_style: Style for body text
        """
        from reportlab.platypus import Paragraph, Spacer

        # Process content
        lines = content.split('\n')

        # Parse all tables from the content first
        tables = self._parse_markdown_table(lines)
        table_line_ranges = []

        # Find line ranges that contain tables
        if tables:
            table_line_ranges = self._find_table_line_ranges(lines)

        # Process lines, skipping table lines and adding tables at
        # appropriate positions
        i = 0
        current_table_idx = 0

        while i < len(lines):
            line = lines[i].strip()

            # Check if this line is part of a table
            is_table_line = any(
                start <= i <= end for start, end in table_line_ranges
            )

            if is_table_line:
                # Skip all lines in this table and add the table to story
                table_start, table_end = next(
                    (start, end)
                    for start, end in table_line_ranges
                    if start <= i <= end
                )

                if current_table_idx < len(tables):
                    try:
                        table = self._create_pdf_table(
                            tables[current_table_idx]
                        )
                        story.append(table)
                        story.append(Spacer(1, 12))
                    except Exception as e:
                        logger.error(f"Failed to create table: {e}")
                        # Fallback: render as text
                        table_error_msg = (
                            f"Table data (error): "
                            f"{tables[current_table_idx]}"
                        )
                        story.append(
                            Paragraph(
                                table_error_msg,
                                body_style,
                            )
                        )
                    current_table_idx += 1

                # Skip to end of table
                i = table_end + 1
                continue

            # Skip empty lines
            if not line:
                story.append(Spacer(1, 6))
                i += 1
                continue

            # Handle headings
            if line.startswith('# '):
                story.append(Paragraph(line[2:], heading_style))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], heading_style))
            elif line.startswith('### '):
                story.append(Paragraph(line[4:], heading_style))
            else:
                # Regular paragraph
                # Convert basic markdown formatting
                line = self._convert_markdown_to_html(line)
                story.append(Paragraph(line, body_style))

            i += 1

    def _find_table_line_ranges(
        self, lines: List[str]
    ) -> List[Tuple[int, int]]:
        r"""Find line ranges that contain markdown tables.

        Args:
            lines (List[str]): List of lines to analyze.

        Returns:
            List[Tuple[int, int]]: List of (start_line, end_line) tuples
                for table ranges.
        """
        ranges = []
        in_table = False
        table_start = 0

        for i, line in enumerate(lines):
            line = line.strip()

            if self._is_table_row(line):
                if not in_table:
                    in_table = True
                    table_start = i
            else:
                if in_table:
                    # End of table
                    ranges.append((table_start, i - 1))
                    in_table = False

        # Handle table at end of content
        if in_table:
            ranges.append((table_start, len(lines) - 1))

        return ranges

    def _register_chinese_font(self) -> str:
        r"""Register Chinese font for PDF generation.

        Returns:
            str: The font name to use for Chinese text.
        """
        import os
        import platform

        from reportlab.lib.fonts import addMapping
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        # Try to find and register Chinese fonts on the system
        font_paths = []
        system = platform.system()

        if system == "Darwin":  # macOS
            font_paths = [
                "/System/Library/Fonts/PingFang.ttc",
                "/System/Library/Fonts/Hiragino Sans GB.ttc",
                "/System/Library/Fonts/STHeiti Light.ttc",
                "/System/Library/Fonts/STHeiti Medium.ttc",
                "/Library/Fonts/Arial Unicode MS.ttf",
            ]
        elif system == "Windows":
            font_paths = [
                r"C:\Windows\Fonts\msyh.ttc",  # Microsoft YaHei
                r"C:\Windows\Fonts\simsun.ttc",  # SimSun
                r"C:\Windows\Fonts\arial.ttf",  # Arial (fallback)
            ]
        elif system == "Linux":
            font_paths = [
                "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            ]

        # Try to register the first available font
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    font_name = "ChineseFont"
                    # Only register if not already registered
                    if font_name not in pdfmetrics.getRegisteredFontNames():
                        pdfmetrics.registerFont(TTFont(font_name, font_path))
                        # Add font mapping for bold/italic variants
                        addMapping(font_name, 0, 0, font_name)  # normal
                        addMapping(font_name, 0, 1, font_name)  # italic
                        addMapping(font_name, 1, 0, font_name)  # bold
                        addMapping(font_name, 1, 1, font_name)  # bold italic
                    return font_name
                except Exception:
                    continue

        # Fallback to Helvetica if no Chinese font found
        logger.warning("No Chinese font found, falling back to Helvetica")
        return "Helvetica"

    def _parse_markdown_table(self, lines: List[str]) -> List[List[List[str]]]:
        r"""Parse markdown-style tables from a list of lines.

        Args:
            lines (List[str]): List of text lines that may contain tables.

        Returns:
            List[List[List[str]]]: List of tables, where each table is a list
                of rows, and each row is a list of cells.
        """
        tables = []
        current_table_data: List[List[str]] = []
        in_table = False

        for line in lines:
            line = line.strip()

            # Check for table (Markdown-style)
            if self._is_table_row(line):
                if not in_table:
                    in_table = True
                    current_table_data = []

                # Skip separator lines (e.g., |---|---|)
                if self._is_table_separator(line):
                    continue

                # Parse table row
                cells = self._parse_table_row(line)
                if cells:
                    current_table_data.append(cells)
                continue

            # If we were in a table and now we're not, finalize the table
            if in_table:
                if current_table_data:
                    tables.append(current_table_data)
                current_table_data = []
                in_table = False

        # Add any remaining table
        if in_table and current_table_data:
            tables.append(current_table_data)

        return tables

    def _is_table_row(self, line: str) -> bool:
        r"""Check if a line appears to be a table row.

        Args:
            line (str): The line to check.

        Returns:
            bool: True if the line looks like a table row.
        """
        return '|' in line and line.count('|') >= 2

    def _is_table_separator(self, line: str) -> bool:
        r"""Check if a line is a table separator (e.g., |---|---|).

        Args:
            line (str): The line to check.

        Returns:
            bool: True if the line is a table separator.
        """
        import re

        # More precise check for separator lines
        # Must contain only spaces, pipes, dashes, and colons
        # and have at least one dash to be a separator
        if not re.match(r'^[\s\|\-\:]+$', line):
            return False

        # Must contain at least one dash to be a valid separator
        return '-' in line

    def _parse_table_row(self, line: str) -> List[str]:
        r"""Parse a single table row into cells.

        Args:
            line (str): The table row line.

        Returns:
            List[str]: List of cell contents.
        """
        # Parse table row
        cells = [cell.strip() for cell in line.split('|')]

        # Remove empty cells at start/end (common in markdown tables)
        if cells and not cells[0]:
            cells = cells[1:]
        if cells and not cells[-1]:
            cells = cells[:-1]

        return cells

    def _create_pdf_table(self, table_data: List[List[str]]):
        r"""Create a formatted table for PDF.

        Args:
            table_data (List[List[str]]): Table data as list of rows.

        Returns:
            Table: A formatted reportlab Table object.
        """
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import Paragraph, Table, TableStyle

        try:
            # Get Chinese font for table
            chinese_font = self._register_chinese_font()

            # Calculate available width (A4 width minus margins)
            page_width = A4[0]  # A4 width in points
            margins = 144  # left + right margins (72 each)
            available_width = page_width - margins

            # Calculate column widths and font size based on content
            if table_data:
                num_columns = len(table_data[0])

                # Calculate max content length for each column
                max_lengths = [0] * num_columns
                max_cell_length = 0
                for row in table_data:
                    for i, cell in enumerate(row):
                        if i < len(max_lengths):
                            cell_length = len(str(cell))
                            max_lengths[i] = max(max_lengths[i], cell_length)
                            max_cell_length = max(max_cell_length, cell_length)

                # Dynamic font size calculation based on columns and content
                # Base font sizes
                base_header_font = 9
                base_body_font = 8

                # Calculate font size factor based on columns and content
                column_factors = {10: 0.6, 8: 0.7, 6: 0.8, 4: 0.9}
                font_size_factor = next(
                    (
                        factor
                        for cols, factor in column_factors.items()
                        if num_columns > cols
                    ),
                    1.0,
                )

                # Further adjust if max cell content is very long
                if max_cell_length > 30:
                    font_size_factor *= 0.8
                elif max_cell_length > 20:
                    font_size_factor *= 0.9

                header_font_size = max(
                    5, int(base_header_font * font_size_factor)
                )
                body_font_size = max(5, int(base_body_font * font_size_factor))

                # Calculate minimum column width based on font size
                min_col_width = max(30, 40 * font_size_factor)

                # Distribute width proportionally with minimum width
                total_length = sum(max_lengths)
                if total_length > 0:
                    # Calculate proportional widths
                    proportional_widths = [
                        (length / total_length) * available_width
                        for length in max_lengths
                    ]

                    # Ensure minimum width and adjust if necessary
                    col_widths = []
                    total_width = 0
                    for width in proportional_widths:
                        adjusted_width = max(min_col_width, width)
                        col_widths.append(adjusted_width)
                        total_width += adjusted_width

                    # Scale down if total exceeds available width
                    if total_width > available_width:
                        scale_factor = available_width / total_width
                        col_widths = [w * scale_factor for w in col_widths]
                else:
                    col_widths = [available_width / num_columns] * num_columns

                # Adjust padding based on font size
                h_padding = max(2, int(6 * font_size_factor))
                v_padding = max(2, int(4 * font_size_factor))
            else:
                col_widths = None
                header_font_size = 9
                body_font_size = 8
                h_padding = 6
                v_padding = 4

            # Create paragraph styles for wrapping text
            header_style = ParagraphStyle(
                'TableHeader',
                fontName=chinese_font,
                fontSize=header_font_size,
                textColor=colors.whitesmoke,
                alignment=0,  # LEFT alignment
                leading=header_font_size * 1.2,
            )

            body_style = ParagraphStyle(
                'TableBody',
                fontName=chinese_font,
                fontSize=body_font_size,
                textColor=colors.black,
                alignment=0,  # LEFT alignment
                leading=body_font_size * 1.2,
            )

            # Convert table data to Paragraph objects for text wrapping
            wrapped_data = []
            for row_idx, row in enumerate(table_data):
                wrapped_row = []
                for cell in row:
                    cell_text = str(cell)
                    # Use header style for first row, body style for others
                    style = header_style if row_idx == 0 else body_style
                    # Escape special characters for XML
                    cell_text = (
                        cell_text.replace('&', '&amp;')
                        .replace('<', '&lt;')
                        .replace('>', '&gt;')
                    )
                    para = Paragraph(cell_text, style)
                    wrapped_row.append(para)
                wrapped_data.append(wrapped_row)

            # Create table with wrapped data
            table = Table(wrapped_data, colWidths=col_widths, repeatRows=1)

            # Style the table with dynamic formatting
            table.setStyle(
                TableStyle(
                    [
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('LEFTPADDING', (0, 0), (-1, -1), h_padding),
                        ('RIGHTPADDING', (0, 0), (-1, -1), h_padding),
                        ('TOPPADDING', (0, 0), (-1, -1), v_padding),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), v_padding),
                    ]
                )
            )

            return table

        except Exception as e:
            logger.error(f"Error creating table: {e}")
            # Return simple unstyled table as fallback
            from reportlab.platypus import Table

            return Table(table_data)

    def _convert_markdown_to_html(self, text: str) -> str:
        r"""Convert basic markdown formatting to HTML for PDF rendering.

        Args:
            text (str): Text with markdown formatting.

        Returns:
            str: Text with HTML formatting.
        """
        # Define conversion patterns
        conversions = [
            (r'\*\*(.*?)\*\*', r'<b>\1</b>'),  # Bold with **
            (r'__(.*?)__', r'<b>\1</b>'),  # Bold with __
            (
                r'(?<!\*)\*(?!\*)(.*?)(?<!\*)\*(?!\*)',
                r'<i>\1</i>',
            ),  # Italic with *
            (r'(?<!_)_(?!_)(.*?)(?<!_)_(?!_)', r'<i>\1</i>'),  # Italic with _
            (r'`(.*?)`', r'<font name="Courier">\1</font>'),  # Inline code
        ]

        for pattern, replacement in conversions:
            text = re.sub(pattern, replacement, text)

        return text

    def _ensure_html_utf8_meta(self, content: str) -> str:
        r"""Ensure HTML content has UTF-8 meta tag.

        Args:
            content (str): The HTML content.

        Returns:
            str: HTML content with UTF-8 meta tag.
        """
        # Check if content already has a charset meta tag
        has_charset = re.search(
            r'<meta[^>]*charset[^>]*>', content, re.IGNORECASE
        )

        # UTF-8 meta tag
        utf8_meta = '<meta charset="utf-8">'

        if has_charset:
            # Replace existing charset with UTF-8
            content = re.sub(
                r'<meta[^>]*charset[^>]*>',
                utf8_meta,
                content,
                flags=re.IGNORECASE,
            )
        else:
            # Add UTF-8 meta tag
            # Try to find <head> tag
            head_match = re.search(r'<head[^>]*>', content, re.IGNORECASE)
            if head_match:
                # Insert after <head> tag
                insert_pos = head_match.end()
                content = (
                    content[:insert_pos]
                    + '\n    '
                    + utf8_meta
                    + content[insert_pos:]
                )
            else:
                # No <head> tag found, check if there's <html> tag
                html_match = re.search(r'<html[^>]*>', content, re.IGNORECASE)
                if html_match:
                    # Insert <head> with meta tag after <html>
                    insert_pos = html_match.end()
                    content = (
                        content[:insert_pos]
                        + '\n<head>\n    '
                        + utf8_meta
                        + '\n</head>'
                        + content[insert_pos:]
                    )
                else:
                    # No proper HTML structure, wrap content
                    content = (
                        '<!DOCTYPE html>\n<html>\n<head>\n    '
                        + utf8_meta
                        + '\n</head>\n<body>\n'
                        + content
                        + '\n</body>\n</html>'
                    )

        return content

    def _write_csv_file(
        self,
        file_path: Path,
        content: Union[str, List[List]],
        encoding: str = "utf-8",
    ) -> None:
        r"""Write CSV content to a file.

        Args:
            file_path (Path): The target file path.
            content (Union[str, List[List]]): The CSV content as a string or
                list of lists.
            encoding (str): Character encoding to use. (default: :obj:`utf-8`)
        """
        import csv

        with file_path.open("w", encoding=encoding, newline='') as f:
            if isinstance(content, str):
                f.write(content)
            else:
                writer = csv.writer(f)
                writer.writerows(content)

    def _write_json_file(
        self,
        file_path: Path,
        content: str,
        encoding: str = "utf-8",
    ) -> None:
        r"""Write JSON content to a file.

        Args:
            file_path (Path): The target file path.
            content (str): The JSON content as a string.
            encoding (str): Character encoding to use. (default: :obj:`utf-8`)
        """
        import json

        with file_path.open("w", encoding=encoding) as f:
            if isinstance(content, str):
                try:
                    # Try parsing as JSON string first
                    data = json.loads(content)
                    json.dump(data, f, ensure_ascii=False)
                except json.JSONDecodeError:
                    # If not valid JSON string, write as is
                    f.write(content)
            else:
                # If not string, dump as JSON
                json.dump(content, f, ensure_ascii=False)

    def _write_simple_text_file(
        self, file_path: Path, content: str, encoding: str = "utf-8"
    ) -> None:
        r"""Write text content to a file (used for HTML, Markdown, YAML, etc.).

        Args:
            file_path (Path): The target file path.
            content (str): The content to write.
            encoding (str): Character encoding to use. (default: :obj:`utf-8`)
        """
        # For HTML files, ensure UTF-8 meta tag is present
        if file_path.suffix.lower() in ['.html', '.htm']:
            content = self._ensure_html_utf8_meta(content)

        with file_path.open("w", encoding=encoding) as f:
            f.write(content)

    def write_to_file(
        self,
        title: str,
        content: Union[str, List[List[str]]],
        filename: str,
        encoding: Optional[str] = None,
        use_latex: bool = False,
    ) -> str:
        r"""Write the given content to a file.

        If the file exists, it will be overwritten. Supports multiple formats:
        Markdown (.md, .markdown, default), Plaintext (.txt), CSV (.csv),
        DOC/DOCX (.doc, .docx), PDF (.pdf), JSON (.json), YAML (.yml, .yaml),
        and HTML (.html, .htm).

        Args:
            title (str): The title of the document.
            content (Union[str, List[List[str]]]): The content to write to the
                file. Content format varies by file type:
                - Text formats (txt, md, html, yaml): string
                - CSV: string or list of lists
                - JSON: string or serializable object
            filename (str): The name or path of the file. If a relative path is
                supplied, it is resolved to self.working_directory.
            encoding (Optional[str]): The character encoding to use. (default:
                :obj: `None`)
            use_latex (bool): Whether to use LaTeX for math rendering.
                (default: :obj:`False`)

        Returns:
            str: A message indicating success or error details.
        """
        file_path = self._resolve_filepath(filename)
        file_path.parent.mkdir(parents=True, exist_ok=True)

        # Generate unique filename if file exists
        file_path = self._generate_unique_filename(file_path)

        extension = file_path.suffix.lower()

        # If no extension is provided, use markdown as default
        if extension == "":
            file_path = file_path.with_suffix('.md')
            extension = '.md'

        try:
            # Get encoding or use default
            file_encoding = encoding or self.default_encoding

            if extension in [".doc", ".docx"]:
                self._write_docx_file(file_path, str(content))
            elif extension == ".pdf":
                self._write_pdf_file(file_path, title, content, use_latex)
            elif extension == ".csv":
                self._write_csv_file(
                    file_path, content, encoding=file_encoding
                )
            elif extension == ".json":
                self._write_json_file(
                    file_path,
                    content,  # type: ignore[arg-type]
                    encoding=file_encoding,
                )
            elif extension in [
                ".yml",
                ".yaml",
                ".html",
                ".htm",
                ".md",
                ".markdown",
            ]:
                self._write_simple_text_file(
                    file_path, str(content), encoding=file_encoding
                )
            else:
                # Fallback to simple text writing for unknown or .txt
                # extensions
                self._write_text_file(
                    file_path, str(content), encoding=file_encoding
                )

            msg = f"Content successfully written to file: {file_path}"
            logger.info(msg)
            return msg
        except Exception as e:
            error_msg = (
                f"Error occurred while writing to file {file_path}: {e}"
            )
            logger.error(error_msg)
            return error_msg

    # ----------------------------------------------
    # Read File Functions
    # ----------------------------------------------
    def read_file(self, file_path: str) -> str:
        r"""Read and return content of a file with automatic format
        detection.

        This method automatically detects the file format based on its
        extension and uses the appropriate reader method. It supports
        multiple file formats including text files, documents, PDFs,
        structured data files, and web files.

        Args:
            file_path (str): The path to the file to read. Can be
                relative or absolute. If relative, it will be resolved
                relative to the working directory.

        Returns:
            str: The content of the file as a string. For structured
                formats like JSON and YAML, the content is formatted with
                proper indentation. For binary formats like PDFs, the
                extracted text content is returned.
        """
        try:
            file_path = self._resolve_filepath(file_path)

            extension = file_path.suffix.lower()
            if extension in [".doc", ".docx"]:
                return self._read_docx_file(file_path)
            elif extension == ".pdf":
                return self._read_pdf_file(file_path)
            elif extension == ".json":
                return self._read_json_file(file_path)
            elif extension in [".yaml", ".yml"]:
                return self._read_yaml_file(file_path)
            else:
                return self._read_text_file(file_path)
        except Exception as e:
            return f"Error reading file: {e}"

    @dependencies_required('docx')
    def _read_docx_file(self, file_path: Path) -> str:
        r"""Read the content of a DOCX file.

        Args:
            file_path (Path): The path to the file to read.
        """

        import docx

        doc = docx.Document(str(file_path))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    @dependencies_required('PyPDF2')
    def _read_pdf_file(self, file_path: Path) -> str:
        r"""Read the content of a PDF file.

        Args:
            file_path (Path): The path to the file to read.
        """
        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            return f"Error reading PDF file: {e}"

    def _read_text_file(self, file_path: Path) -> str:
        r"""Read the content of a text file.

        Args:
            file_path (Path): The path to the file to read.
        """
        with file_path.open("r", encoding=self.default_encoding) as f:
            return f.read()

    def _read_json_file(self, file_path: Path) -> str:
        r"""Read the content of a JSON file and return as formatted string.

        Args:
            file_path (Path): The path to the file to read.
        """
        import json

        try:
            with file_path.open("r", encoding=self.default_encoding) as f:
                data = json.load(f)
                return json.dumps(data, indent=2, ensure_ascii=False)
        except Exception as e:
            return f"Error reading JSON file: {e}"

    @dependencies_required('yaml')
    def _read_yaml_file(self, file_path: Path) -> str:
        r"""Read the content of a YAML file and return as formatted string.

        Args:
            file_path (Path): The path to the file to read.
        """
        import yaml

        try:
            with file_path.open("r", encoding=self.default_encoding) as f:
                data = yaml.safe_load(f)
                return yaml.dump(
                    data, default_flow_style=False, allow_unicode=True
                )
        except Exception as e:
            return f"Error reading YAML file: {e}"

    # ----------------------------------------------
    # Edit File Functions
    # ----------------------------------------------
    def edit_file(
        self, file_path: str, old_content: str, new_content: str
    ) -> str:
        r"""Edit the content of a file.

        This method provides file editing capabilities that automatically
        detect the file format and apply appropriate editing strategies. It
        supports both content replacement and structured editing for
        different file types.

        Args:
            file_path (str): The path to the file to edit. Can be
                relative or absolute. If relative, it will be resolved
                relative to the working directory. If the file doesn't
                exist, it will be created.
            old_content (str): The content to replace. For text files,
                this is a literal string to find and replace. For
                structured files like JSON, this can be a JSON path
                (e.g., "$.key.subkey") or a key to modify.
            new_content (str): The new content to write. For text files,
                this replaces the old content. For structured files, this
                is the new value to set.

        Returns:
            str: A success message if the edit was successful, or an
                error message containing details about what went wrong.
        """
        try:
            working_path = self._resolve_filepath(file_path)

            # Determine file type and edit accordingly
            extension = working_path.suffix.lower()

            if extension in [".doc", ".docx"]:
                result = self._edit_docx_file(
                    working_path, old_content, new_content
                )
            elif extension == ".pdf":
                result = self._edit_pdf_file(
                    working_path, old_content, new_content
                )
            elif extension == ".json":
                result = self._edit_json_file(
                    working_path, old_content, new_content
                )
            elif extension in [".yaml", ".yml"]:
                result = self._edit_yaml_file(
                    working_path, old_content, new_content
                )
            else:
                result = self._text_replace(
                    working_path, old_content, new_content
                )

            return result

        except Exception as e:
            return f"Error editing file: {e}"

    def _edit_json_file(
        self, file_path: Path, old_content: str, new_content: str
    ) -> str:
        r"""Edit a JSON file by replacing content.

        Args:
            file_path (Path): The path to the JSON file to edit.
            old_content (str): The content to replace (can be a JSON
                path or key).
            new_content (str): The new content to add.
        """
        import json

        try:
            # Read existing JSON data
            with file_path.open("r", encoding=self.default_encoding) as f:
                data = json.load(f)

            # Try to parse old_content as JSON path or key
            if old_content.startswith("$.") or "." in old_content:
                # Handle JSON path-like replacement
                keys = old_content.replace("$.", "").split(".")
                current = data
                for key in keys[:-1]:
                    if isinstance(current, dict) and key in current:
                        current = current[key]
                    else:
                        return f"Path {old_content} not found in JSON"

                # Try to parse new_content as JSON, fallback to string
                try:
                    new_data = json.loads(new_content)
                except json.JSONDecodeError:
                    new_data = new_content

                current[keys[-1]] = new_data

            # Write back to file
            with file_path.open("w", encoding=self.default_encoding) as f:
                json.dump(data, f, indent=2, ensure_ascii=False)
            return f"JSON file {file_path} edited successfully."
        except Exception as e:
            logger.error(f"Error editing JSON file: {e}")
            return f"Error editing JSON file: {e}"

    @dependencies_required('PyYAML')
    def _edit_yaml_file(
        self, file_path: Path, old_content: str, new_content: str
    ) -> str:
        r"""Edit a YAML file by replacing content.

        Args:
            file_path (Path): The path to the YAML file to edit.
            old_content (str): The content to replace (can be a YAML
                path or key).
            new_content (str): The new content to add.
        """
        try:
            import yaml

            # Read existing YAML data
            with file_path.open("r", encoding=self.default_encoding) as f:
                data = yaml.safe_load(f)

            # Try to parse old_content as YAML path or key
            if "." in old_content:
                # Handle YAML path-like replacement
                keys = old_content.split(".")
                current = data
                for key in keys[:-1]:
                    if isinstance(current, dict) and key in current:
                        current = current[key]
                    else:
                        return f"Path {old_content} not found in YAML"

                # Try to parse new_content as YAML, fallback to string
                try:
                    new_data = yaml.safe_load(new_content)
                except yaml.YAMLError:
                    new_data = new_content

                current[keys[-1]] = new_data
            else:
                # Simple string replacement in the entire YAML
                yaml_str = yaml.dump(
                    data, default_flow_style=False, allow_unicode=True
                )
                if old_content in yaml_str:
                    updated_str = yaml_str.replace(old_content, new_content)
                    data = yaml.safe_load(updated_str)
                else:
                    # If old_content not found, append to root level
                    try:
                        new_data = yaml.safe_load(new_content)
                        if isinstance(data, dict):
                            data.update(new_data)
                        else:
                            data = new_data
                    except yaml.YAMLError:
                        # If new_content is not valid YAML, add as a new key
                        if isinstance(data, dict):
                            data["new_content"] = new_content

            # Write back to file
            with file_path.open("w", encoding=self.default_encoding) as f:
                yaml.dump(
                    data, f, default_flow_style=False, allow_unicode=True
                )

            return f"YAML file {file_path} edited successfully."
        except Exception as e:
            logger.error(f"Error editing YAML file: {e}")
            return f"Error editing YAML file: {e}"

    @dependencies_required('docx')
    def _edit_docx_file(
        self, file_path: Path, old_content: str, new_content: str
    ) -> str:
        r"""Edit a DOCX file by replacing content.

        Args:
            file_path (Path): The path to the DOCX file to edit.
            old_content (str): The content to replace.
            new_content (str): The new content to add.
        """
        import docx

        try:
            doc = docx.Document(str(file_path))
            # Check if we're replacing the entire document content
            # Read the current document content
            current_content = "\n".join(
                [paragraph.text for paragraph in doc.paragraphs]
            )

            if old_content.strip() == current_content.strip():
                # Replace entire document content
                # Clear existing paragraphs
                for paragraph in doc.paragraphs[:]:
                    p = paragraph._element
                    p.getparent().remove(p)

                # Add new content as paragraphs
                for line in new_content.split('\n'):
                    if line.strip():  # Only add non-empty lines
                        doc.add_paragraph(line.strip())
            else:
                # Try to replace content within individual paragraphs
                for paragraph in doc.paragraphs:
                    if old_content in paragraph.text:
                        paragraph.text = paragraph.text.replace(
                            old_content, new_content
                        )

            doc.save(str(file_path))
            return f"DOCX file {file_path} edited successfully."
        except Exception as e:
            logger.error(f"Error editing DOCX file: {e}")
            return f"Error editing DOCX file: {e}"

    def _edit_pdf_file(
        self, file_path: Path, old_content: str, new_content: str
    ) -> str:
        r"""Edit a PDF file by creating a new version with modified content.

        Since PDFs are complex binary files, this function creates a new PDF
        with the modified content rather than editing the existing one.

        Args:
            file_path (Path): The path to the PDF file to edit.
            old_content (str): The content to replace.
            new_content (str): The new content to add.
        """
        try:
            # Read the original PDF content
            original_text = self._read_pdf_file(file_path)

            # Replace the old content with new content
            if old_content in original_text:
                modified_text = original_text.replace(old_content, new_content)
            else:
                # If old content not found, append new content
                modified_text = original_text + "\n\n" + new_content

            # Create a new PDF with the modified content
            self._write_pdf_file(
                file_path,
                title="Modified PDF Document",
                content=modified_text,
                use_latex=False,
            )
            return f"PDF file {file_path} edited successfully."
        except Exception as e:
            logger.error(f"Error editing PDF file: {e}")
            return f"Error editing PDF file: {e}"

    def _text_replace(
        self, file_path: Path, old_content: str, new_content: str
    ) -> str:
        r"""Replace specified string in the file.
        Use for updating specific content in the file.

        Args:
            file_path (Path): The path to the file to edit.
            old_content (str): The old content to be replaced.
            new_content (str): The new content to be replaced.

        Returns:
            str: A message indicating the result of the operation.
        """
        try:
            file_text = file_path.read_text(encoding=self.default_encoding)
            old_content = old_content.expandtabs()
            new_content = new_content.expandtabs()

            occurrences = file_text.count(old_content)
            if occurrences == 0:
                return (
                    f"No replacement performed: '{old_content}' not found in "
                    f"{file_path}."
                )
            elif occurrences > 1:
                file_text_line = file_text.split('\n')
                lines = [
                    idx + 1
                    for idx, line in enumerate(file_text_line)
                    if old_content in line
                ]
                return (
                    f"Multiple occurrences of '{old_content}' found in "
                    f"{file_path} at lines {lines}. "
                    "Please specify the line number to replace."
                )

            new_file_content = file_text.replace(old_content, new_content)
            file_path.write_text(
                new_file_content, encoding=self.default_encoding
            )

            return f"Content replaced in {file_path} successfully."
        except Exception as e:
            return f"Error replacing content in {file_path}: {e}"

    def get_tools(self) -> List[FunctionTool]:
        r"""Return a list of FunctionTool objects representing the functions
        in the toolkit.

        Returns:
            List[FunctionTool]: A list of FunctionTool objects representing
                the available functions in this toolkit.
        """
        return [
            FunctionTool(self.write_to_file),
            FunctionTool(self.read_file),
            FunctionTool(self.edit_file),
        ]
