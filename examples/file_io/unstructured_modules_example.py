# =========== Copyright 2023 @ CAMEL-AI.org. All Rights Reserved. ===========
# Licensed under the Apache License, Version 2.0 (the “License”);
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an “AS IS” BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
# =========== Copyright 2023 @ CAMEL-AI.org. All Rights Reserved. ===========

import os

import docx

from camel.file_io.unstructured_io import UnstructuredModules

unstructured_modules = UnstructuredModules()


def parse_file_example():
    document = docx.Document()
    document.add_paragraph("Important Analysis", style="Heading 1")
    document.add_paragraph("Here is my first thought.", style="Body Text")
    document.add_paragraph("Here is my second thought.", style="Normal")
    document.save("mydoc.docx")
    elements = unstructured_modules.parse_file_or_url("mydoc.docx")
    # Cleanup: remove the created file after the example
    if os.path.exists("mydoc.docx"):
        os.remove("mydoc.docx")
    return [" ".join(str(el).split()) for el in elements]


def parse_url_example():
    url = ("https://www.cnn.com/2023/01/30/sport/empire-state-building-green-"
           "philadelphia-eagles-spt-intl/index.html")
    elements = unstructured_modules.parse_file_or_url(url)
    return "\n\n".join([str(el) for el in elements])


def clean_text_example():
    example_text = "Some dirty text â€™ with extra spaces and – dashes."
    example_options = {
        'replace_unicode_quotes': {},
        'clean_extra_whitespace': {},
        'clean_dashes': {},
        'clean_non_ascii_chars': {}
    }
    return unstructured_modules.clean_text_data(clean_options=example_options,
                                                text=example_text)


def extract_data_example():
    example_email_text = "Contact me at example@email.com."
    return unstructured_modules.extract_data_from_text(
        extract_type="extract_email_address", text=example_email_text)


def stage_data_example():
    example_url = (
        "https://www.cnn.com/2023/01/30/sport/empire-state-building-green-"
        "philadelphia-eagles-spt-intl/index.html")
    example_elements = unstructured_modules.parse_file_or_url(example_url)

    staged_element = unstructured_modules.stage_elements(
        stage_type="stage_for_baseplate", elements=example_elements)
    return staged_element


def chunk_url_content_example():
    exampe_url = (
        "https://www.cnn.com/2023/01/30/sport/empire-state-building-green-"
        "philadelphia-eagles-spt-intl/index.html")
    example_elements = unstructured_modules.parse_file_or_url(exampe_url)
    chunks = unstructured_modules.chunk_elements(chunk_type="chunk_by_title",
                                                 elements=example_elements)
    return chunks


def main():

    print("Choose an example to run:")
    print("1. Parse File")
    print("2. Parse URL")
    print("3. Clean Text")
    print("4. Extract Data")
    print("5. Stage Data")
    print("6. Chunk URL Content")
    choice = input("Enter your choice (1-6): ")

    if choice == '1':
        print("Parsing file example:")
        print(parse_file_example())

    elif choice == '2':
        print("Parsing URL example:")
        print(parse_url_example())
        print("\n You have parsed https://www.cnn.com/2023/01/30/sport/"
              "empire-state-building-green-philadelphia-eagles-spt-"
              "intl/index.html")

    elif choice == '3':
        print("Cleaning text example:")
        print("\n Some dirty text â€™ with extra spaces and – dashes."
              "\n \n into \n")
        print(clean_text_example())

    elif choice == '4':
        print("Extracting email example:")
        print(extract_data_example())
        print("extracted from")
        print("Contact me at example@email.com.")

    elif choice == '5':
        print("Staging data example:")
        print(stage_data_example())

    elif choice == '6':
        print("Chunking URL content example:")
        for chunk in chunk_url_content_example():
            print(chunk)
            print("\n" + "-" * 80)

    else:
        print("Invalid choice.")


if __name__ == "__main__":
    main()
